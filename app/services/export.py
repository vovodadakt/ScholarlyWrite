from __future__ import annotations
import io
import re as _re

from docx import Document
from docx.shared import Inches, Pt
from fpdf import FPDF

from app.services.references import (
    resolve_citations, build_bibliography_map, format_reference, STYLE_META,
)


def _load_figures(project_id: str) -> list[dict]:
    """Load figure metadata for a project. Returns list of dicts with storage_path, caption, etc."""
    import os as _os
    from pathlib import Path as _Path
    figures = []
    root = _Path(__file__).parent.parent
    from app.database import async_session
    from sqlalchemy import select as _select
    import asyncio as _asyncio

    try:
        loop = _asyncio.get_event_loop()
        if loop.is_running():
            return figures
    except RuntimeError:
        pass

    return figures


def _get_figures_sync(project_id: str) -> list[dict]:
    """Load figures for a project synchronously via a helper DB query."""
    import os as _os
    from pathlib import Path as _Path
    figures = []
    figures_dir = _Path(__file__).parent.parent / "static" / "figures"
    if not figures_dir.exists():
        return figures

    for f in sorted(figures_dir.iterdir()):
        if f.suffix.lower() in {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}:
            figures.append({
                "storage_path": f"figures/{f.name}",
                "caption": "",
                "figure_number": 0,
            })
    return figures


def build_docx(title: str, chapters: list[dict], references: list[dict] | None = None, figures: list[dict] | None = None, journal_style: str = "apa") -> io.BytesIO:
    from pathlib import Path as _Path
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # Resolve citations in all chapter content
    refs_list = references or []
    for ch in chapters:
        ch["content"] = resolve_citations(ch["content"], refs_list, journal_style)

    heading = doc.add_heading(title, level=0)
    heading.alignment = 1

    for ch in chapters:
        doc.add_heading(ch["title"], level=1)
        for paragraph in ch["content"].split("\n"):
            paragraph = paragraph.strip()
            if paragraph:
                fig_match = _re.match(r'^\[图\s*(\d+)\](.*)', paragraph)
                if fig_match and figures:
                    fig_num = int(fig_match.group(1))
                    rest = fig_match.group(2).strip()
                    for f in figures:
                        if f.get("figure_number") == fig_num:
                            doc.add_paragraph(f"Figure {fig_num}: {f.get('caption', '')}")
                            img_path = _Path(__file__).parent.parent / "static" / f["storage_path"]
                            if img_path.exists():
                                try:
                                    doc.add_picture(str(img_path), width=Inches(5))
                                except Exception:
                                    pass
                            if rest:
                                doc.add_paragraph(rest)
                            break
                    else:
                        doc.add_paragraph(paragraph)
                else:
                    doc.add_paragraph(paragraph)

    if refs_list:
        bib_map = build_bibliography_map(refs_list, journal_style)
        ref_section_title = STYLE_META.get(journal_style, {}).get("name", "参考文献")
        doc.add_heading("参考文献", level=1)
        for key, entry in bib_map.items():
            doc.add_paragraph(entry, style="List Number")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_latex(title: str, chapters: list[dict], references: list[dict] | None = None, figures: list[dict] | None = None, journal_style: str = "apa") -> str:
    lines = [
        r"\documentclass[12pt,a4paper]{article}",
        r"\usepackage[UTF8]{ctex}",
        r"\usepackage{geometry}",
        r"\geometry{margin=2.54cm}",
        r"\usepackage{hyperref}",
        r"\usepackage{amsmath,amssymb}",
        r"\usepackage{graphicx}",
        r"\title{" + title + "}",
        r"\author{}",
        r"\date{}",
        r"\begin{document}",
        r"\maketitle",
    ]

    # Resolve citations
    refs_list = references or []
    for ch in chapters:
        ch["content"] = resolve_citations(ch["content"], refs_list, journal_style)

    for ch in chapters:
        title_clean = ch["title"].replace("&", r"\&").replace("#", r"\#").replace("_", r"\_")
        lines.append(r"\section{" + title_clean + "}")
        for para in ch["content"].split("\n"):
            para = para.strip()
            if para:
                fig_match = _re.match(r'^\[图\s*(\d+)\](.*)', para)
                if fig_match and figures:
                    fig_num = int(fig_match.group(1))
                    rest = fig_match.group(2).strip()
                    matched = None
                    for f in figures:
                        if f.get("figure_number") == fig_num:
                            matched = f
                            break
                    if matched:
                        cap = matched.get("caption", "").replace("&", r"\&").replace("#", r"\#").replace("_", r"\_").replace("%", r"\%")
                        img_path = "figures/" + matched.get("storage_path", "").split("/")[-1]
                        lines.append(r"\begin{figure}[htbp]")
                        lines.append(r"\centering")
                        lines.append(r"\includegraphics[width=0.8\textwidth]{" + img_path + "}")
                        lines.append(r"\caption{图" + str(fig_num) + ": " + cap + "}")
                        lines.append(r"\end{figure}")
                        if rest:
                            para_clean = rest.replace("&", r"\&").replace("#", r"\#").replace("_", r"\_").replace("%", r"\%")
                            lines.append(para_clean)
                            lines.append("")
                    else:
                        para_clean = para.replace("&", r"\&").replace("#", r"\#").replace("_", r"\_").replace("%", r"\%")
                        lines.append(para_clean)
                        lines.append("")
                else:
                    para_clean = para.replace("&", r"\&").replace("#", r"\#").replace("_", r"\_").replace("%", r"\%")
                    lines.append(para_clean)
                    lines.append("")

    if refs_list:
        lines.append(r"\section*{参考文献}")
        lines.append(r"\begin{enumerate}")
        bib_map = build_bibliography_map(refs_list, journal_style)
        for key, entry in bib_map.items():
            entry_clean = entry.replace("&", r"\&").replace("#", r"\#").replace("_", r"\_").replace("%", r"\%")
            lines.append(r"\item " + entry_clean)
        lines.append(r"\end{enumerate}")

    lines.append(r"\end{document}")
    return "\n".join(lines)


def _format_single_ref_docx(ref: dict) -> str:
    authors = ref.get("authors", [])
    if isinstance(authors, str):
        authors = authors.split(", ")
    year = ref.get("year", "")
    title = ref.get("title", "")
    journal = ref.get("journal", "")
    doi = ref.get("doi", "")

    parts = []
    if authors:
        parts.append("; ".join(authors))
    if year:
        parts.append(f"({year})")
    parts.append(title)
    if journal:
        parts.append(journal)
    if doi:
        parts.append(f"doi:{doi}")
    return ". ".join(parts) + "."


def _format_single_ref_latex(ref: dict) -> str:
    return _format_single_ref_docx(ref)


def _find_cjk_font() -> tuple[str | None, str | None]:
    """Find a CJK-capable font on the system. Returns (regular_path, bold_path)."""
    from pathlib import Path
    import platform

    candidates = []
    if platform.system() == "Windows":
        candidates = [
            ("C:/Windows/Fonts/msyh.ttc", "C:/Windows/Fonts/msyhbd.ttc"),
            ("C:/Windows/Fonts/simsun.ttc", "C:/Windows/Fonts/simhei.ttf"),
            ("C:/Windows/Fonts/NotoSansSC-Regular.ttf", "C:/Windows/Fonts/NotoSansSC-Bold.ttf"),
        ]
    else:
        candidates = [
            ("/usr/share/fonts/truetype/noto/NotoSansSC-Regular.ttf",
             "/usr/share/fonts/truetype/noto/NotoSansSC-Bold.ttf"),
            ("/usr/share/fonts/opentype/noto/NotoSansSC-Regular.otf",
             "/usr/share/fonts/opentype/noto/NotoSansSC-Bold.otf"),
            ("/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc",
             "/usr/share/fonts/noto-cjk/NotoSansCJK-Bold.ttc"),
        ]

    # Also check bundled fonts
    fonts_dir = Path(__file__).parent.parent / "static" / "fonts"
    candidates.append((str(fonts_dir / "NotoSansSC-Regular.ttf"), str(fonts_dir / "NotoSansSC-Bold.ttf")))

    for reg, bold in candidates:
        if Path(reg).exists():
            return reg, bold if Path(bold).exists() else reg
    return None, None


def build_pdf(title: str, chapters: list[dict], references: list[dict] | None = None, figures: list[dict] | None = None, journal_style: str = "apa") -> io.BytesIO:
    from pathlib import Path as _Path

    # Resolve citations
    refs_list = references or []
    for ch in chapters:
        ch["content"] = resolve_citations(ch["content"], refs_list, journal_style)

    pdf = FPDF()
    pdf.add_page()

    font_regular, font_bold = _find_cjk_font()
    use_cjk = font_regular is not None and font_bold is not None
    if use_cjk:
        pdf.add_font("CJK", "", font_regular)
        pdf.add_font("CJK", "B", font_bold)
        body_font = "CJK"
    else:
        body_font = "Helvetica"

    # Title
    pdf.set_font(body_font, "B", 18)
    pdf.multi_cell(0, 12, title, align="C")
    pdf.ln(8)

    for ch in chapters:
        pdf.set_font(body_font, "B", 14)
        pdf.multi_cell(0, 10, ch["title"])
        pdf.ln(2)

        pdf.set_font(body_font, "", 11)
        for para in ch["content"].split("\n"):
            para = para.strip()
            if para:
                # Check for figure placeholder [图1] or [图 1]
                fig_match = _re.match(r'^\[图\s*(\d+)\](.*)', para)
                if fig_match and figures:
                    fig_num = int(fig_match.group(1))
                    rest = fig_match.group(2).strip()
                    matched = None
                    for f in figures:
                        if f.get("figure_number") == fig_num:
                            matched = f
                            break
                    if matched:
                        img_path = _Path(__file__).parent.parent / "static" / matched["storage_path"]
                        caption = matched.get("caption", "")
                        if img_path.exists():
                            try:
                                pdf.set_font(body_font, "", 10)
                                pdf.multi_cell(0, 6, f"图{fig_num}: {caption}", align="C")
                                pdf.ln(2)
                                # Insert image, fit to page width
                                pdf.image(str(img_path), x=pdf.get_x() + 10, w=pdf.w - 20)
                                pdf.ln(4)
                            except Exception:
                                pdf.multi_cell(0, 7, para)
                        else:
                            pdf.multi_cell(0, 7, para)
                        if rest:
                            pdf.set_font(body_font, "", 11)
                            pdf.multi_cell(0, 7, rest)
                            pdf.ln(2)
                    else:
                        pdf.multi_cell(0, 7, para)
                        pdf.ln(2)
                else:
                    pdf.multi_cell(0, 7, para)
                    pdf.ln(2)

    if refs_list:
        pdf.add_page()
        pdf.set_font(body_font, "B", 14)
        pdf.multi_cell(0, 10, "参考文献")
        pdf.ln(4)
        pdf.set_font(body_font, "", 10)
        bib_map = build_bibliography_map(refs_list, journal_style)
        for key, entry in bib_map.items():
            pdf.multi_cell(0, 6, entry)
            pdf.ln(2)

    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf


def build_cover_letter(project_title: str, authors: str, journal_name: str,
                       editor_name: str = "", highlights: list[str] | None = None,
                       custom_message: str = "") -> io.BytesIO:
    """Generate a cover letter DOCX for journal submission."""
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # Sender info
    date_str = ""
    try:
        from datetime import date
        date_str = date.today().strftime("%B %d, %Y")
    except Exception:
        pass

    if date_str:
        doc.add_paragraph(date_str)

    doc.add_paragraph("")
    greeting = f"Dear Editor"
    if editor_name:
        greeting = f"Dear Dr. {editor_name}"
    elif journal_name:
        greeting = f"Dear Editor of {journal_name}"
    doc.add_paragraph(greeting)
    doc.add_paragraph("")

    # Body
    body = doc.add_paragraph()
    body.add_run(
        f"I am pleased to submit our manuscript entitled "
        f"\"{project_title}\" for consideration for publication in {journal_name}. "
    )

    if custom_message:
        body.add_run(custom_message)
    else:
        body.add_run(
            "Our research presents novel findings that we believe will be of significant "
            "interest to the readership of your journal. The work represents a substantial "
            "contribution to the field, employing rigorous methodology and comprehensive analysis."
        )

    doc.add_paragraph("")

    # Highlights
    if highlights:
        doc.add_paragraph("Key highlights of our study include:")
        for h in highlights:
            doc.add_paragraph(h, style="List Bullet")
        doc.add_paragraph("")

    # Declarations
    doc.add_paragraph("We confirm that:")
    declarations = [
        "This manuscript has not been published elsewhere and is not under consideration by another journal.",
        "All authors have approved the manuscript and agree with its submission to " + journal_name + ".",
        "The authors declare no conflicts of interest.",
        "All sources of funding have been acknowledged in the manuscript.",
    ]
    for d in declarations:
        doc.add_paragraph(d, style="List Bullet")

    doc.add_paragraph("")

    # Closing
    doc.add_paragraph("We appreciate your time and consideration, and look forward to your response.")
    doc.add_paragraph("")
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph("")
    if authors:
        doc.add_paragraph(authors)
    doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
